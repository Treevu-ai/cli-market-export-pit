"""Convert executive proposal markdown to a styled Word document."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

GREEN = RGBColor(0x0F, 0x3D, 0x2E)


def _clean_lines(md_text: str) -> list[str]:
    lines = md_text.splitlines()
    cleaned: list[str] = []
    for line in lines:
        if line.strip() == "" and cleaned and cleaned[-1] == "":
            continue
        cleaned.append(line.rstrip())
    while cleaned and cleaned[0] == "":
        cleaned.pop(0)
    while cleaned and cleaned[-1] == "":
        cleaned.pop()
    return cleaned


def _add_rich_paragraph(doc: Document, text: str, *, style: str | None = None, italic: bool = False):
    paragraph = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
            if italic:
                run.italic = True
    return paragraph


def _parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("|"):
            break
        if re.match(r"^\|[-:\s|]+\|$", line):
            i += 1
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            value = row[c_idx] if c_idx < len(row) else ""
            value = re.sub(r"\*\*([^*]+)\*\*", r"\1", value)
            value = re.sub(r"\*([^*]+)\*", r"\1", value)
            table.rows[r_idx].cells[c_idx].text = value
    doc.add_paragraph()


def export_docx(input_md: Path, output_docx: Path, base_dir: Path) -> None:
    lines = _clean_lines(input_md.read_text(encoding="utf-8"))
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped == "":
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            heading = doc.add_heading(stripped[2:].strip(), level=0)
            for run in heading.runs:
                run.font.color.rgb = GREEN
            i += 1
            continue

        if stripped.startswith("## "):
            heading = doc.add_heading(stripped[3:].strip(), level=1)
            for run in heading.runs:
                run.font.color.rgb = GREEN
            i += 1
            continue

        if stripped.startswith("### "):
            heading = doc.add_heading(stripped[4:].strip(), level=2)
            for run in heading.runs:
                run.font.color.rgb = GREEN
            i += 1
            continue

        if stripped.startswith("```"):
            code_lines: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            continue

        if stripped.startswith("|"):
            rows, i = _parse_table(lines, i)
            _add_table(doc, rows)
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            alt, src = image_match.groups()
            image_path = (base_dir / src).resolve()
            if image_path.exists():
                doc.add_paragraph(alt, style="Intense Quote")
                doc.add_picture(str(image_path), width=Inches(6.2))
                doc.add_paragraph()
            else:
                _add_rich_paragraph(doc, f"[Imagen: {alt}]")
            i += 1
            continue

        if stripped.startswith("> "):
            quote_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("> "):
                quote_lines.append(lines[i].strip()[2:].strip())
                i += 1
            paragraph = doc.add_paragraph("\n".join(quote_lines), style="Intense Quote")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if re.match(r"^\d+\.\s", stripped):
            items: list[str] = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i].strip()):
                items.append(re.sub(r"^\d+\.\s*", "", lines[i].strip()))
                i += 1
            for item in items:
                _add_rich_paragraph(doc, item, style="List Number")
            continue

        if stripped.startswith("- "):
            items = []
            while i < len(lines) and lines[i].strip().startswith("- "):
                items.append(lines[i].strip()[2:].strip())
                i += 1
            for item in items:
                _add_rich_paragraph(doc, item, style="List Bullet")
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            _add_rich_paragraph(doc, stripped.strip("*"), italic=True)
            i += 1
            continue

        _add_rich_paragraph(doc, stripped)
        i += 1

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def try_save_as_doc(docx_path: Path, doc_path: Path) -> bool:
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(docx_path.resolve()))
        document.SaveAs(str(doc_path.resolve()), FileFormat=0)  # wdFormatDocument
        document.Close(False)
        word.Quit()
        return True
    except Exception:
        return False


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_md", type=Path)
    parser.add_argument("-o", "--output", type=Path)
    parser.add_argument("--also-doc", action="store_true")
    args = parser.parse_args()
    output = args.output or args.input_md.with_suffix(".docx")
    export_docx(args.input_md, output, args.input_md.parent.resolve())
    print(output)
    if args.also_doc or output.suffix.lower() == ".doc":
        doc_path = output.with_suffix(".doc")
        if try_save_as_doc(output if output.suffix.lower() == ".docx" else output.with_suffix(".docx"), doc_path):
            print(doc_path)


if __name__ == "__main__":
    main()
